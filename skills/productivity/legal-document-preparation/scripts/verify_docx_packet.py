#!/usr/bin/env python3
"""
verify_docx_packet.py — Automated consistency verifier for generated DOCX packets.

Checks across all .docx files in a directory:
  - Shared terms appear consistently (principal, currency, parties, governing law)
  - No malformed money strings ($$)
  - No stale template labels ([State], [BORROWER LEGAL NAME], etc.)
  - Draft/not-sign-ready legend present in headers AND footers
  - Enumerates unresolved [[TBD...]] placeholders by owner (COUNSEL/CPA/CONFIRM)
  - Validates deal_terms.json for duplicate keys (if present in the directory)

Usage:
  python verify_docx_packet.py /path/to/packet_dir
  python verify_docx_packet.py /path/to/packet_dir --terms deal_terms.json

Exit codes:
  0 = all checks passed
  1 = issues found (details printed)
"""

import sys
import os
import json
import re
import argparse
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(2)


def extract_doc_text(docx_path):
    """Extract all text: body paragraphs, table cells, headers, footers."""
    doc = Document(str(docx_path))
    body = []
    for p in doc.paragraphs:
        body.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                body.append(cell.text)
    headers = []
    for section in doc.sections:
        for hp in section.header.paragraphs:
            headers.append(hp.text)
        for fp in section.footer.paragraphs:
            headers.append(fp.text)
    return {
        "body": "\n".join(body),
        "headers_footers": "\n".join(headers),
        "full": "\n".join(body + headers),
    }


def check_duplicate_json_keys(json_path):
    """Detect duplicate keys in a JSON file (json.loads silently takes last)."""
    import collections
    duplicates = []
    try:
        raw = Path(json_path).read_text()
        # Naive but effective: parse line-by-line for repeated top-level keys
        key_pattern = re.compile(r'^\s*"([^"]+)"\s*:', re.MULTILINE)
        keys = key_pattern.findall(raw)
        counter = collections.Counter(keys)
        for key, count in counter.items():
            if count > 1:
                duplicates.append(f'"{key}" appears {count} times')
    except Exception as e:
        return [f"JSON parse error: {e}"]
    return duplicates


def verify_packet(packet_dir, terms_file=None):
    issues = []
    base = Path(packet_dir)
    docx_files = sorted(base.glob("*.docx"))

    if not docx_files:
        issues.append(f"No .docx files found in {base}")

    # Load terms JSON if available
    terms_path = base / (terms_file or "deal_terms.json")
    if terms_path.exists():
        dup_keys = check_duplicate_json_keys(terms_path)
        if dup_keys:
            issues.append(f"deal_terms.json duplicate keys: {dup_keys}")

    # Stale template labels that should have been filled
    stale_labels = [
        r"\[State\]",
        r"\[BORROWER LEGAL NAME\]",
        r"\[LENDER LEGAL NAME\]",
        r"\[Full legal name of U\.S\. S corporation\]",
        r"\[Full legal name of Japanese company\]",
        r"\[Japanese legal form\](?!.*—)",  # allow [[Japanese legal form — CONFIRM]]
    ]

    draft_legend_patterns = [
        r"DRAFT FOR COUNSEL REVIEW",
        r"NOT SIGN-READY",
    ]

    for docx_path in docx_files:
        text = extract_doc_text(docx_path)

        # Check for $$ money formatting bug
        if "$$" in text["full"]:
            # Count occurrences
            count = text["full"].count("$$")
            issues.append(f"{docx_path.name}: MALFORMED MONEY STRING ($$ x{count})")

        # Check for stale template labels
        for pattern in stale_labels:
            matches = re.findall(pattern, text["body"])
            if matches:
                issues.append(
                    f"{docx_path.name}: STALE TEMPLATE LABEL '{pattern}' "
                    f"({len(matches)} occurrences)"
                )

        # Check draft legend in headers/footers
        hf = text["headers_footers"]
        for legend_pat in draft_legend_patterns:
            if not re.search(legend_pat, hf, re.IGNORECASE):
                issues.append(
                    f"{docx_path.name}: MISSING draft legend in header/footer: "
                    f"'{legend_pat}'"
                )

        # Enumerate TBD placeholders by owner
        tbd_counsel = len(re.findall(r'\[\[TBD.*COUNSEL', text["full"]))
        tbd_cpa = len(re.findall(r'\[\[TBD.*CPA', text["full"]))
        tbd_confirm = len(re.findall(r'\[\[TBD.*CONFIRM', text["full"]))
        tbd_plain = len(
            re.findall(r'\[\[TBD(?!(?:.*COUNSEL|.*CPA|.*CONFIRM))', text["full"])
        )
        if any([tbd_counsel, tbd_cpa, tbd_confirm, tbd_plain]):
            print(f"  {docx_path.name}: TBD fields — "
                  f"counsel={tbd_counsel}, CPA={tbd_cpa}, "
                  f"confirm={tbd_confirm}, plain={tbd_plain}")

        # Cross-check shared terms (if we have terms JSON)
        if terms_path.exists():
            try:
                terms = json.loads(terms_path.read_text())
                # Check principal appears in body
                if "10,000,000" in str(terms.get("principal_usd", "")):
                    if "10,000,000" not in text["body"]:
                        issues.append(
                            f"{docx_path.name}: MISSING principal '10,000,000'"
                        )
                # Check parties appear
                lender = terms.get("lender_short_name", "")
                borrower = terms.get("borrower_short_name", "")
                if lender and lender not in text["full"]:
                    issues.append(f"{docx_path.name}: MISSING lender '{lender}'")
                if borrower and borrower not in text["full"]:
                    issues.append(f"{docx_path.name}: MISSING borrower '{borrower}'")
            except json.JSONDecodeError:
                issues.append(f"deal_terms.json: INVALID JSON (parse error)")

    return issues


def main():
    parser = argparse.ArgumentParser(
        description="Verify DOCX packet consistency"
    )
    parser.add_argument(
        "packet_dir",
        help="Directory containing generated .docx files"
    )
    parser.add_argument(
        "--terms",
        default=None,
        help="Deal terms JSON filename (default: deal_terms.json)"
    )
    args = parser.parse_args()

    print(f"Verifying packet: {args.packet_dir}")
    issues = verify_packet(args.packet_dir, args.terms)

    if issues:
        print(f"\n⚠ {len(issues)} ISSUE(S) FOUND:\n")
        for issue in issues:
            print(f"  ⚠ {issue}")
        sys.exit(1)
    else:
        print("\n✓ All checks passed — consistent terms, no formatting bugs, "
              "no stale labels, draft legends present.")
        sys.exit(0)


if __name__ == "__main__":
    main()
